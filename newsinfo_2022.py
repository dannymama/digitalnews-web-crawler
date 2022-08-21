
#For tiny url
from __future__ import with_statement
import contextlib
from urllib.parse import urlencode
from urllib.request import urlopen
import sys
import time

def make_tiny(url):
    request_url = ('http://tinyurl.com/api-create.php?' + urlencode({'url':url}))
    with contextlib.closing(urlopen(request_url)) as response:
        return response.read().decode('utf-8')
def main():
    for tinyurl in map(make_tiny, sys.argv[1:]):
        print(tinyurl)
if __name__ == '__main__':
    main()

#For web crawler; pip install "requests" & "bs4" required in advance
import requests
import time
from bs4 import BeautifulSoup

from urllib.request import urlretrieve

#For output as docx & pdf; pip install "python-docx" & "comtypes" required in advance
from docx import Document
from docx.shared import Inches

import sys
import os
import comtypes.client

#For transforming image format
from PIL import Image

#Initialize/customize news search criteria
keywords = []
paths = ['']
num = 100
start = 0
while True:
    option = input('Use default keywords(數位金融,行動支付,第三方支付,金融科技) input 0\nDefine keywords on your own input 1\n')
    if option =='1':
        keywords.append(input('Input your keywords(use OR to append multiple): '))
        break
    elif option == '0':
        keywords = ['數位金融OR行動支付OR第三方支付OR金融科技']
        break
    elif option == 'num':
        num = int(input('input num: '))
    elif option == 'start':
        start = int(input('input start: '))
    else:
        print('Please input again ')

print('Loading...')

hdr = {'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.11 (KHTML, like Gecko) Chrome/23.0.1271.64 Safari/537.11',
       'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
       'Accept-Charset': 'ISO-8859-1,utf-8;q=0.7,*;q=0.3',
       'Accept-Encoding': 'none',
       'Accept-Language': 'en-US,en;q=0.8',
       'Connection': 'keep-alive'}


#Google News url, with q=defined keywords & cr=countryTW & num=100, no pagination
res = requests.get('https://www.google.com.tw/search?q=%s&source=lnms&tbm=nws&tbs=qdr:w&cr=countryTW&num=%d&start=%d'%(keywords[0],num,start))
soup = BeautifulSoup(res.text, 'lxml')

newscontent = []

#To crawl TW popular news sites

ctime = {} #工商時報
chn = {} #中時新聞網
bnext = {} #數位時代
eco = {} #經濟日報
technews = {} #科技新報


contentall = {} #Nested in nested structure
timesa = time.strftime("%M")
    

for i in soup.findAll('div', {'class':'Gx5Zad fP1Qef xpd EtOod pkphOe'}):

    try:

        if '中時新聞網' in i.find('div', {'class':'BNeawe UPmit AP7Wnd'}).text:
            title = i.find('div', {'class':'BNeawe vvjwJb AP7Wnd'}).text
            chn[title] = ['','','',''] #[url,contents,source,imageurl]
            url = 'https://www.google.com.tw' + i.a['href']
            chn[title][0] = url #append url
            chn[title][2] = i.find('div', {'class':'BNeawe UPmit AP7Wnd'}).text #append source

            res = requests.get('https://www.google.com.tw' + i.a['href'], timeout =15, headers=hdr)
            res.encoding = 'utf-8'
            soup2 = BeautifulSoup(res.text, 'lxml')
            soup2.encoding ='utf-8'

            for j in soup2.findAll('div', {'class':'article-body'}):
                
                for k in j.findAll('p'):
                    contents = k.text
                    chn[title][1] += contents #append contents
            
            m = soup2.find('div', {'class': 'photo-container'})
            
            try:
                if 'data-src' in str(m.img):
                    chn[title][3] += m.img['data-src'] #append image url
                    try:
                        urlretrieve(chn[title][3], 'C:\\Users\\01539\\Desktop\\diginews\\img\\%s.jpg'%(title)) #download image
                    except:
                        pass
                elif 'src' in str(m.img):
                    chn[title][3] += m.img['src']
                    try:
                        urlretrieve(chn[title][3], 'C:\\Users\\01539\\Desktop\\diginews\\img\\%s.jpg'%(title)) 
                    except:
                        pass
                else:
                    pass
            except:
                pass
            #print(chn[titlechn][3])
    except:
        pass

    try:

        if '數位時代' in i.find('div', {'class':'BNeawe UPmit AP7Wnd'}).text:
            title = i.find('div', {'class':'BNeawe vvjwJb AP7Wnd'}).text
            bnext[title] = ['','','',''] #[url,contents,source,imageurl]
            url = 'https://www.google.com.tw' + i.a['href']
            bnext[title][0] = url #append url
            bnext[title][2] = i.find('div', {'class':'BNeawe UPmit AP7Wnd'}).text #append source

            res = requests.get('https://www.google.com.tw' + i.a['href'], timeout =15, headers=hdr)
            res.encoding = 'utf-8'
            soup2 = BeautifulSoup(res.text, 'lxml')
            soup2.encoding ='utf-8'

            for j in soup2.findAll('div', {'class':'DynamicComp static htmlview'}):
                
                for k in j.findAll('p'):
                    contents = k.text
                    bnext[title][1] += '\n' + contents #append contents
            
            m = soup2.find('div', {'class': 'bigg'})
            
            try:
                if 'data-src' in str(m.img):
                    bnext[title][3] += m.img['data-src'] #append image url
                    try:
                        urlretrieve(bnext[title][3], 'C:\\Users\\01539\\Desktop\\diginews\\img\\%s.jpg'%(title)) #download image
                    except:
                        pass
                elif 'src' in str(m.img):
                    bnext[title][3] += m.img['src']
                    try:
                        urlretrieve(bnext[title][3], 'C:\\Users\\01539\\Desktop\\diginews\\img\\%s.jpg'%(title)) 
                    except:
                        pass
                else:
                    pass
            except:
                pass
    except:
        pass

    try:

        if '科技新報' in i.find('div', {'class':'BNeawe UPmit AP7Wnd'}).text:
            title = i.find('div', {'class':'BNeawe vvjwJb AP7Wnd'}).text
            technews[title] = ['','','',''] #[url,contents,source,imageurl]
            url = 'https://www.google.com.tw' + i.a['href']
            technews[title][0] = url #append url
            technews[title][2] = i.find('div', {'class':'BNeawe UPmit AP7Wnd'}).text #append source

            res = requests.get('https://www.google.com.tw' + i.a['href'], timeout =15, headers=hdr)
            res.encoding = 'utf-8'
            soup2 = BeautifulSoup(res.text, 'lxml')
            soup2.encoding ='utf-8'

            for j in soup2.findAll('div', {'class':'indent'}):
                
                for k in j.findAll('p'):
                    contents = k.text
                    technews[title][1] += contents #append contents
            
            m = soup2.find('div', {'class': 'bigg'})
            
            try:
                if 'data-src' in str(m.img):
                    technews[title][3] += m.img['data-src'] #append image url
                    try:
                        urlretrieve(technews[title][3], 'C:\\Users\\01539\\Desktop\\diginews\\img\\%s.jpg'%(title)) #download image
                    except:
                        pass
                elif 'src' in str(m.img):
                    technews[title][3] += m.img['src']
                    try:
                        urlretrieve(technews[title][3], 'C:\\Users\\01539\\Desktop\\diginews\\img\\%s.jpg'%(title)) 
                    except:
                        pass
                else:
                    pass
            except:
                pass
    except:
        pass


    try:

        if '經濟日報' in i.find('div', {'class':'BNeawe UPmit AP7Wnd'}).text:
            title = i.find('div', {'class':'BNeawe vvjwJb AP7Wnd'}).text
            eco[title] = ['','','',''] #[url,contents,source,imageurl]
            url = 'https://www.google.com.tw' + i.a['href']
            eco[title][0] = url #append url
            eco[title][2] = i.find('div', {'class':'BNeawe UPmit AP7Wnd'}).text #append source

            res = requests.get('https://www.google.com.tw' + i.a['href'], timeout =15, headers=hdr)
            res.encoding = 'utf-8'
            soup2 = BeautifulSoup(res.text, 'lxml')
            soup2.encoding ='utf-8'

            for j in soup2.findAll('section', {'class':'article-body__editor'}):
                
                for k in j.findAll('p'):
                    contents = k.text
                    eco[title][1] += contents #append contents
            
            m = soup2.find('section', {'class': 'article-body__editor'})
            
            try:
                if 'data-src' in str(m.img):
                    eco[title][3] += m.img['data-src'] #append image url
                    try:
                        urlretrieve(eco[title][3], 'C:\\Users\\01539\\Desktop\\diginews\\img\\%s.jpg'%(title)) #download image
                    except:
                        pass
                elif 'src' in str(m.img):
                    eco[title][3] += m.img['src']
                    try:
                        urlretrieve(eco[title][3], 'C:\\Users\\01539\\Desktop\\diginews\\img\\%s.jpg'%(title)) 
                    except:
                        pass
                else:
                    pass
            except:
                pass
    except:
        pass

    try:

        if '工商時報' in i.find('div', {'class':'BNeawe UPmit AP7Wnd'}).text:
            title = i.find('div', {'class':'BNeawe vvjwJb AP7Wnd'}).text
            ctime[title] = ['','','',''] #[url,contents,source,imageurl]
            url = 'https://www.google.com.tw' + i.a['href']
            ctime[title][0] = url #append url
            ctime[title][2] = i.find('div', {'class':'BNeawe UPmit AP7Wnd'}).text #append source

            res = requests.get('https://www.google.com.tw' + i.a['href'], timeout =15, headers=hdr)
            res.encoding = 'utf-8'
            soup2 = BeautifulSoup(res.text, 'lxml')
            soup2.encoding ='utf-8'

            for j in soup2.findAll('div', {'class':'entry-content clearfix single-post-content'}):
                
                for k in j.findAll('p'):
                    contents = k.text
                    ctime[title][1] += '\n' + contents #append contents
            
            m = soup2.find('div', {'class': 'photo-container'})
            
            try:
                if 'data-src' in str(m.img):
                    ctime[title][3] += m.img['data-src'] #append image url
                    try:
                        urlretrieve(ctime[title][3], 'C:\\Users\\01539\\Desktop\\diginews\\img\\%s.jpg'%(title)) #download image
                    except:
                        pass
                elif 'src' in str(m.img):
                    ctime[title][3] += m.img['src']
                    try:
                        urlretrieve(ctime[title][3], 'C:\\Users\\01539\\Desktop\\diginews\\img\\%s.jpg'%(title)) 
                    except:
                        pass
                else:
                    pass
            except:
                pass
    except:
        pass
    

    contentall = {'工商時報':ctime,'中時新聞網':chn,'數位時代':bnext,'經濟日報':eco,'科技新報':technews}

    
#Document docx
times = time.strftime("%m-%d-at %H%M") 
document = Document('C:\\Users\\01539\\Desktop\\diginews\\default.docx') ##create a default docx
document.add_heading('KEYWORDS: %s'%(keywords[0]))
for j in contentall:
    for k in contentall[j]:
        document.add_paragraph(k) #Title list

for j in contentall:
    document.save('C:\\Users\\01539\\Desktop\\diginews\\news_%s.docx'%(times))    
    print('★' +j, ': %d articles'%(len(contentall[j])))
    
    document.add_heading('★' +j)
    for k in contentall[j]:
        document.save('C:\\Users\\01539\\Desktop\\diginews\\news_%s.docx'%(times))
        document.add_heading(k) #Title
        document.add_paragraph(contentall[j][k][2]) #Time
        contentall[j][k][0] = make_tiny(contentall[j][k][0])
        time.sleep(1)

        document.add_paragraph(contentall[j][k][0]) #URL

        try:
            document.add_picture('C:\\Users\\01539\\Desktop\\diginews\\img\\%s.jpg'%(k), height = Inches(2))
        except:
            pass

        document.add_paragraph(contentall[j][k][1]) #URL
        document.add_paragraph('====================================')


document.save('C:\\Users\\01539\\Desktop\\diginews\\news_%s.docx'%(times))

# Save As PDF
wdFormatPDF = 17
word = comtypes.client.CreateObject('Word.Application')
doc = word.Documents.Open('C:\\Users\\01539\\Desktop\\diginews\\news_%s.docx'%(times))
doc.SaveAs('C:\\Users\\01539\\Desktop\\diginews\\news_%s.pdf'%(times), FileFormat=wdFormatPDF)
doc.Close()
word.Quit()


print('\nCompleted...\n File path:\n', 'C:\\Users\\01539\\Desktop\\diginews\\news_%s.pdf'%(times))
timesb = time.strftime("%M")

long = int(timesb) - int(timesa)
print('\nRun time: %dmins'%(long))
print(input('\nInput 0 to finish: '))
