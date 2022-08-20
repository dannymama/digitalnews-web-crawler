#D:\\diginews\\

#以下for縮網址
from __future__ import with_statement
import contextlib
from urllib.parse import urlencode
from urllib.request import urlopen
import sys
import time
'''print(sys.stdout.encoding) #cp950, 中文的 windows「命令提示字元」(cmd) 編碼預設'''
def make_tiny(url):
    request_url = ('http://tinyurl.com/api-create.php?' + urlencode({'url':url}))
    with contextlib.closing(urlopen(request_url)) as response:
        return response.read().decode('utf-8')
def main():
    for tinyurl in map(make_tiny, sys.argv[1:]):
        print(tinyurl)
if __name__ == '__main__':
    main()

#以下for爬蟲
import requests
import time
from bs4 import BeautifulSoup

from urllib.request import urlretrieve

#以下for輸出docx pip install python-docx
from docx import Document
from docx.shared import Inches

#For圖片轉換格式
from PIL import Image


keywords = []
paths = ['']
num = 100
start = 0
while True:
    option = input('使用預設關鍵字(數位金融,行動支付,第三方支付,金融科技)請輸入0\n自行輸入關鍵字請輸入1\n')
    if option =='1':
        keywords.append(input('請輸入搜尋關鍵字: '))
        break
    elif option == '0':
        keywords = ['數位金融OR行動支付OR第三方支付OR金融科技']
        break
    elif option == 'num':
        num = int(input('input num: '))
    elif option == 'start':
        start = int(input('input start: '))
    else:
        print('請重新輸入 ')

print('Loading...')

hdr = {'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.11 (KHTML, like Gecko) Chrome/23.0.1271.64 Safari/537.11',
       'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
       'Accept-Charset': 'ISO-8859-1,utf-8;q=0.7,*;q=0.3',
       'Accept-Encoding': 'none',
       'Accept-Language': 'en-US,en;q=0.8',
       'Connection': 'keep-alive'}

#加一個輸入1bankID來自動下載到桌面

#取得google news搜尋網址, 無設計翻頁, 一次就是100則新聞 先10篇, 可更改start
res = requests.get('https://www.google.com.tw/search?q=%s&source=lnms&tbm=nws&tbs=qdr:w&cr=countryTW&num=%d&start=%d'%(keywords[0],num,start))
soup = BeautifulSoup(res.text, 'lxml')
'''.encoding ='utf-8'''

newscontent = []

udn1 = {} #聯合財經網
udn2 = {} #聯合新聞網
chn = {} #中時電子報
apple = {} #蘋果日報
moneydj = {} #MoneyDJ理財網
etnews = {} #ETNEWS
bnext = {} #數位時代
ithome = {} #iThome
technews = {} #科技新報
ltn = {} #自由時報
lnl = {} #關鍵評論

contentall = {} #雙層朝狀結構
timesa = time.strftime("%M")
#找title, url
for i in soup.findAll('h3', {'class':'r'}):
    fullurl = 'https://www.google.com.tw' + i.a['href']
    '''tinyurl = make_tiny(fullurl)''' #先取消tinyurl
    '''newscontent.append({'title':i.text})''' ###
    '''newscontent.append({'url':fullurl})''' ###

#找from 包含時間資訊
for i in soup.findAll('div', {'class':'slp'}):
    
    if '數位時代' in i.span.text:
        titlebnext = i.find_previous_sibling().text #標題
        bnext[titlebnext] = ['','','','']
        urlbnext = 'https://www.google.com.tw' + i.find_previous_sibling().a['href'] #連結
        bnext[titlebnext][0] = urlbnext
        bnext[titlebnext][2] = i.span.text

        resbnext = requests.get('https://www.google.com.tw' + i.find_previous_sibling().a['href'], timeout =15, headers=hdr)
        soupbnext = BeautifulSoup(resbnext.text, 'lxml')
        soupbnext.encoding = 'utf-8'

        for j in soupbnext.findAll('article', {'class':'main_content'}):
            for k in j.findAll('p'):
                contentbnext = k.text #內容
                bnext[titlebnext][1] += '\n' + contentbnext #字串連接! 非常好用
                
        '''if soupbnext.findAll('div', {'top_pic'}): #讚讚'''
        styletext = str(soupbnext.select('div > #top_pic')[0])
        styleurl = styletext.partition("url('")[-1].partition("');")[0]
        bnext[titlebnext][3] += styleurl
        try:
            urlretrieve(bnext[titlebnext][3], 'D:\\diginews\\img\\%s.jpg'%(titlebnext)) #下載圖片
        except:
            print(titlebnext +' 圖片下載失敗')
        
    '''if '聯合財經網' in i.span.text:

        titleudn1 = i.find_previous_sibling().text #標題
        udn1[titleudn1] = ['','','','']
        urludn1 = 'https://www.google.com.tw' + i.find_previous_sibling().a['href'] #連結
        udn1[titleudn1][0] = urludn1
        udn1[titleudn1][2] = i.span.text

        resbudn1 = requests.get('https://www.google.com.tw' + i.find_previous_sibling().a['href'], timeout =15)
        
        soupudn1 = BeautifulSoup(resudn1.text, 'lxml')
        soupudn1.encoding ='utf-8'

        for j in soupudn1.findAll('div', {'id':'story_body_content'}):
            for k in j.findAll('p'):
                contentudn1 = k.text #內容
                udn1[titleudn1][1] += '\n' + contentudn1 #字串連接! 非常好用'''
        
    if '聯合新聞網' in i.span.text:
        titleudn2 = i.find_previous_sibling().text #標題
        udn2[titleudn2] = ['','','','']
        urludn2 = 'https://www.google.com.tw' + i.find_previous_sibling().a['href'] #連結
        udn2[titleudn2][0] = urludn2
        udn2[titleudn2][2] = i.span.text

        resudn2 = requests.get('https://www.google.com.tw' + i.find_previous_sibling().a['href'], timeout =15, headers=hdr)
        resudn2.encoding = 'utf-8'
        soupudn2 = BeautifulSoup(resudn2.text, 'lxml')
        soupudn2.encoding ='utf-8'

        for j in soupudn2.findAll('div', {'id':'story_body_content'}):
            for k in j.findAll('p'):
                contentudn2 = k.text #內容
                udn2[titleudn2][1] += '\n' + contentudn2

        m = soupudn2.find('div', {'id': 'story_body_content'}) #讚讚
        try:
            if 'data-src' in str(m.img):
                udn2[titleudn2][3] += m.img['data-src']
                try:
                    urlretrieve(udn2[titleudn2][3], 'D:\\diginews\\img\\%s.jpg'%(titleudn2)) #下載圖片
                except:
                    print(titleudn2 +' 圖片下載失敗')
            elif 'src' in str(m.img):
                udn2[titleudn2][3] += m.img['src']
                try:
                    urlretrieve(udn2[titleudn2][3], 'D:\\diginews\\img\\%s.jpg'%(titleudn2)) #下載圖片
                except:
                    print(titleudn2 +' 圖片下載失敗')
            else:
                pass
        except:
            pass
                
        
    if '中時電子報' in i.span.text:
        titlechn = i.find_previous_sibling().text #標題
        chn[titlechn] = ['','','','','']
        urlchn = 'https://www.google.com.tw' + i.find_previous_sibling().a['href'] #連結
        chn[titlechn][0] = urlchn
        chn[titlechn][2] = i.span.text

        reschn = requests.get('https://www.google.com.tw' + i.find_previous_sibling().a['href'], timeout =15, headers=hdr)
        soupchn = BeautifulSoup(reschn.text, 'lxml')
        soupchn.encoding ='utf-8'

        for j in soupchn.findAll('article', {'class':'clear-fix'}):
            if bool(j.p): ##bool
                for k in j.findAll('p'):
                    contentchn = k.text #內容
                    chn[titlechn][1] += '\n' + contentchn #字串連接! 非常好用

        for m in soupchn.findAll('div', {'class':'img_view'}): #讚讚
            try:
                chn[titlechn][3] += m.find('img')['src']
                urlretrieve(chn[titlechn][3], 'D:\\diginews\\img\\%s.jpg'%(titlechn)) #下載圖片
            except:
                pass

    '''        
    if '蘋果日報' in i.span.text:
        titleapple = i.find_previous_sibling().text #標題
        apple[titleapple] = ['','','']
        urlapple = 'https://www.google.com.tw' + i.find_previous_sibling().a['href'] #連結
        apple[titleapple][0] = urlapple
        apple[titleapple][2] = i.span.text

        resapple = requests.get('https://www.google.com.tw' + i.find_previous_sibling().a['href'])
        soupapple = BeautifulSoup(resapple.text, 'lxml')

        for j in soupapple.findAll('p', {'id':'summary'}):
            contentapple = j.text #內容
            apple[titleapple][1] += contentapple
            '''
            
    if 'MoneyDJ理財網' in i.span.text:
        titlemoneydj = i.find_previous_sibling().text #標題
        moneydj[titlemoneydj] = ['','','','']
        urlmoneydj = 'https://www.google.com.tw' + i.find_previous_sibling().a['href'] #連結
        moneydj[titlemoneydj][0] = urlmoneydj
        moneydj[titlemoneydj][2] = i.span.text

        resmoneydj = requests.get('https://www.google.com.tw' + i.find_previous_sibling().a['href'], timeout =15, headers=hdr)
        soupmoneydj = BeautifulSoup(resmoneydj.text, 'lxml')
        soupmoneydj.encoding ='utf-8'

        for j in soupmoneydj.findAll('div', {'class':'wikilink'}):
            for k in j.findAll('p'):
                contentmoneydj = k.text #內容
                moneydj[titlemoneydj][1] += '\n' + contentmoneydj #字串連接! 非常好用
            
    if 'ETNEWS' in i.span.text:
        titleetnews = i.find_previous_sibling().text #標題
        etnews[titleetnews] = ['','','','']
        urletnews = 'https://www.google.com.tw' + i.find_previous_sibling().a['href'] #連結
        etnews[titleetnews][0] = urletnews
        etnews[titleetnews][2] = i.span.text

        resetnews = requests.get('https://www.google.com.tw' + i.find_previous_sibling().a['href'], timeout =15, headers=hdr)
        soupetnews = BeautifulSoup(resetnews.text, 'lxml')
        soupetnews.encoding ='utf-8'

        for j in soupetnews.findAll('div', {'class':'story'}):
            for k in j.findAll('p'):
                contentetnews = k.text #內容
                etnews[titleetnews][1] += '\n' + contentetnews #字串連接! 非常好用

        m = soupetnews.find('div', {'class':'story'})
        try:
            if 'src' in str(m.img):
                etnews[titleetnews][3] += 'http:' + m.img['src']

                try:
                    '''urlretrieve(lnl[titlelnl][3], 'D:\\diginews\\img\\%s.jpg'%(titlelnl)) #下載圖片'''
                    with open('D:\\diginews\\img\\%s.jpg'%(titleetnews), 'wb') as f:
                        res = requests.get(etnews[titleetnews][3], headers=hdr)
                        f.write(res.content)
                except:
                    pass
        except:
            pass
            
    if 'iThome' in i.span.text:
        titleithome = i.find_previous_sibling().text #標題
        ithome[titleithome] = ['','','','']
        urlithome = 'https://www.google.com.tw' + i.find_previous_sibling().a['href'] #連結
        ithome[titleithome][0] = urlithome
        ithome[titleithome][2] = i.span.text

        resithome = requests.get('https://www.google.com.tw' + i.find_previous_sibling().a['href'], timeout =15, headers=hdr)
        soupithome = BeautifulSoup(resithome.text, 'lxml')
        soupithome.encoding ='utf-8'

        for j in soupithome.findAll('div', {'class':'contents-wrap'}):
            for k in j.findAll('p'):
                contentithome = k.text #內容
                ithome[titleithome][1] += '\n' + contentithome #字串連接! 非常好用

        m = soupithome.find('div', {'class':'field-item'})
        try:
            if 'src' in str(m.img):
                ithome[titleithome][3] += m.img['src'].partition('?')[0]

                '''try:'''
                '''urlretrieve(lnl[titlelnl][3], 'D:\\diginews\\img\\titlelnl%s.123'%(titlelnl)) #下載圖片'''
                with open('D:\\diginews\\img\\%s.123'%(titleithome), 'wb') as f:
                    res = requests.get(ithome[titleithome][3], headers=hdr)
                    f.write(res.content)
                g =Image.open('D:\\diginews\\img\\%s.123'%(titleithome)) ###先設為.123
                g.save('D:\\diginews\\img\\%s.jpg'%(titleithome)) ###再轉為.jpg就可以
                
                '''except:
                    pass'''
        except:
            pass
        
    if '科技新報' in i.span.text:
        titletechnews = i.find_previous_sibling().text #標題
        technews[titletechnews] = ['','','','']
        urltechnews = 'https://www.google.com.tw' + i.find_previous_sibling().a['href'] #連結
        technews[titletechnews][0] = urltechnews
        technews[titletechnews][2] = i.span.text

        restechnews = requests.get('https://www.google.com.tw' + i.find_previous_sibling().a['href'], timeout =15, headers=hdr)
        souptechnews = BeautifulSoup(restechnews.text, 'lxml')
        souptechnews.encoding ='utf-8'

        for j in souptechnews.findAll('div', {'class':'indent'}):
            for k in j.findAll('p'):
                contenttechnews = k.text #內容
                technews[titletechnews][1] += '\n' + contenttechnews #字串連接! 非常好用

        m = souptechnews.find('img', {'class':'attachment-post-thumbnail'})
        try:
            if 'src' in str(m):
                technews[titletechnews][3] += m['src']

                try:
                    '''urlretrieve(lnl[titlelnl][3], 'D:\\diginews\\img\\%s.jpg'%(titlelnl)) #下載圖片'''
                    with open('D:\\diginews\\img\\%s.jpg'%(titletechnews), 'wb') as f:
                        res = requests.get(technews[titletechnews][3], headers=hdr)
                        f.write(res.content)
                except:
                    pass
        except:
            pass

            
    if '自由時報' in i.span.text:
        titleltn = i.find_previous_sibling().text #標題
        ltn[titleltn] = ['','','','']
        urlltn = 'https://www.google.com.tw' + i.find_previous_sibling().a['href'] #連結
        ltn[titleltn][0] = urlltn
        ltn[titleltn][2] = i.span.text

        resltn = requests.get('https://www.google.com.tw' + i.find_previous_sibling().a['href'], timeout =15, headers=hdr)
        soupltn = BeautifulSoup(resltn.text, 'lxml')
        soupltn.encoding ='utf-8'

        for j in soupltn.findAll('div', {'class':'text'}):
            for k in j.findAll('p'):
                contentltn = k.text #內容
                ltn[titleltn][1] += '\n' + contentltn #字串連接! 非常好用

        m = soupltn.find('img', {'class':'fenixbox'})
        try:
            if 'src' in str(m):
                ltn[titleltn][3] += m['src']

                try:
                    '''urlretrieve(lnl[titlelnl][3], 'D:\\diginews\\img\\%s.jpg'%(titlelnl)) #下載圖片'''
                    with open('D:\\diginews\\img\\%s.jpg'%(titleltn), 'wb') as f:
                        res = requests.get(ltn[titleltn][3], headers=hdr)
                        f.write(res.content)
                except:
                    pass
        except:
            pass
            
    if '關鍵評論' in i.span.text:
        titlelnl = i.find_previous_sibling().text #標題
        lnl[titlelnl] = ['','','','']
        urllnl = 'https://www.google.com.tw' + i.find_previous_sibling().a['href'] #連結
        lnl[titlelnl][0] = urllnl
        lnl[titlelnl][2] = i.span.text

        reslnl = requests.get('https://www.google.com.tw' + i.find_previous_sibling().a['href'], timeout =15, headers=hdr)
        souplnl = BeautifulSoup(reslnl.text, 'lxml')
        souplnl.encoding ='utf-8'

        for j in souplnl.findAll('div', {'class':'article-content'}):
            for k in j.findAll('p'):
                contentlnl = k.text #內容
                lnl[titlelnl][1] += '\n' + contentlnl #字串連接! 非常好用
        m = souplnl.find('img', {'class':'front-img'})
        try:
            if 'src-sm' in str(m):
                lnl[titlelnl][3] += m['src-sm'].partition('?')[0]

                try:
                    '''urlretrieve(lnl[titlelnl][3], 'D:\\diginews\\img\\%s.jpg'%(titlelnl)) #下載圖片'''
                    with open('D:\\diginews\\img\\%s.jpg'%(titlelnl), 'wb') as f:
                        res = requests.get(lnl[titlelnl][3], headers=hdr)
                        f.write(res.content)
                except:
                    pass
        except:
            pass

#先刪除蘋果, 聯合財金
    contentall = {'數位時代':bnext,'聯合新聞網':udn2,'中時電子報':chn,'MoneyDJ理財網':moneydj,'ETNEWS':etnews,'iThome':ithome,'科技新報':technews,'自由時報':ltn,'關鍵評論':lnl}

    
    
'''for j in newscontent:
    print(j)'''

'''times = time.strftime("%m月%d日%H點%M分")'''
'''file2 = open('D:\\Python36-32\\dannyma\\diginews\\news_%s.csv'%(times),'w',encoding='big5')'''
'''file2 = open('D:\\Python36-32\\dannyma\\diginews\\news_%s.txt'%(times),'w',encoding='big5')'''

times = time.strftime("%m-%d-at %H%M") ###
document = Document('D:\\diginews\\default.docx') ###要課製新增空白檔案
document.add_heading('KEYWORDS: %s'%(keywords[0]))
for j in contentall:
    for k in contentall[j]:
        document.add_paragraph(k) #Title 列表

for j in contentall:
    document.save('D:\\diginews\\news_%s.docx'%(times))    
    print('★' +j, ': %d篇'%(len(contentall[j])))
    '''print('====================================', sep =' ', file = file2)'''
    document.add_paragraph('====================================')
    '''print('★' +j, file = file2)'''
    document.add_heading('★' +j)
    for k in contentall[j]:
        document.save('D:\\diginews\\news_%s.docx'%(times))
        '''print('Title: ', k, sep =' ', file = file2)'''
        document.add_heading(k) #Title
        '''print('Time: ', contentall[j][k][2], sep =' ', file = file2)'''
        document.add_paragraph(contentall[j][k][2]) #Time
        contentall[j][k][0] = make_tiny(contentall[j][k][0])
        time.sleep(1)
        '''print('URL: ', contentall[j][k][0], sep =' ', file = file2)'''
        document.add_paragraph(contentall[j][k][0]) #URL

        try:
            document.add_picture('D:\\diginews\\img\\%s.jpg'%(k), height = Inches(2))
        except:
            pass
            '''pass'''
        '''print('Content:\n', contentall[j][k][1], sep =' ', file = file2)'''
        document.add_paragraph(contentall[j][k][1]) #URL
        '''print('===================', sep =' ', file = file2)'''
        document.add_paragraph('===================')

        '''document.add_paragraph(contentall[j][k][3])''' #此為圖片連結



'''file2.close()'''
document.save('D:\\diginews\\news_%s.docx'%(times))
print('\nCompleted...\n File path:\n', 'D:\\diginews\\news_%s.docx'%(times))
timesb = time.strftime("%M")

long = int(timesb) - int(timesa)
print('\n費時: %d分鐘'%(long))
print(input('\n結束請按0: '))
